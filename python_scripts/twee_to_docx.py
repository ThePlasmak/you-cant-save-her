from docx import Document
import re

twee_filename = "src/Base.tw"
word_filename = "export/You_Can't_Save_Her.docx"


def twee_to_word(twee_filename, word_filename):
    # Create a new Document
    doc = Document()

    with open(twee_filename, "r", encoding="utf-8") as file:
        content = file.read()

    # Split the content into passages based on the "::" indicating a new passage
    passages = re.split(r"\n::\s*", content)

    for passage in passages:
        if passage.strip() == "":
            continue  # Skip empty passages
        # Split the passage into title and text based on the first newline
        parts = passage.split("\n", 1)
        title = parts[0].strip()
        text = parts[1].strip() if len(parts) > 1 else ""

        title = title.replace("::", "").strip()

        # Add title as Header 1
        doc.add_heading(title, level=1)

        # Add text as a normal paragraph
        doc.add_paragraph(text)

    # Save the document
    doc.save(word_filename)


twee_to_word(twee_filename, word_filename)
print(f"DOCX exported successfully.")
